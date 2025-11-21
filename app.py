import streamlit as st
from pathlib import Path
from docx import Document  # to read Word (.docx) files
import re                    # clause regex
import pandas as pd          # displaying clause table

# Try to import pdfplumber, but don't crash if it's not installed
try:
    import pdfplumber        # for reading PDF tenders
    PDF_SUPPORT = True
except ImportError:
    pdfplumber = None
    PDF_SUPPORT = False

# Try to import OpenAI client, but don't crash if it's not installed
try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OpenAI = None
    OPENAI_AVAILABLE = False

# ---------- BASIC PAGE SETTINGS (MUST BE FIRST STREAMLIT CALL) ----------
st.set_page_config(
    page_title="Wavetec Tender Library",
    layout="wide"
)

# ---------- SIMPLE & ROBUST LOGIN ----------
def check_password():
    """Returns True if the user entered the correct credentials."""

    # Read auth secrets safely to avoid KeyError
    auth = st.secrets.get("auth", {})
    correct_username = auth.get("username")
    correct_password = auth.get("password")

    # If secrets are missing, show a clear error instead of crashing
    if not correct_username or not correct_password:
        st.error(
            "🔐 Authentication is not configured.\n\n"
            "Please add the following to this app's **Secrets** in Streamlit Cloud:\n\n"
            "[auth]\n"
            'username = "YOUR_USERNAME"\n'
            'password = "YOUR_PASSWORD"'
        )
        return False

    def password_entered():
        """Verify username and password, update session state."""
        entered_username = st.session_state.get("username", "")
        entered_password = st.session_state.get("password", "")

        if entered_username == correct_username and entered_password == correct_password:
            st.session_state["password_correct"] = True
            # Don't keep password in memory
            st.session_state.pop("password", None)
        else:
            st.session_state["password_correct"] = False

    # First run: show login form
    if "password_correct" not in st.session_state:
        st.text_input("Username", key="username")
        st.text_input("Password", type="password", key="password", on_change=password_entered)
        return False

    # If wrong credentials were entered
    if not st.session_state["password_correct"]:
        st.text_input("Username", key="username")
        st.text_input("Password", type="password", key="password", on_change=password_entered)
        st.error("❌ Incorrect username or password")
        return False

    # Correct credentials
    return True


# ---------- STOP APP IF NOT LOGGED IN ----------
if not check_password():
    st.stop()

# ---------- TITLE & INTRO ----------
st.title("📚 Wavetec Tender Library")
st.write("Central knowledge base for automated tender and RFP responses.")

# ---------- PATHS ----------
BASE_DIR = Path(__file__).parent

# Root folder that contains all the library categories
LIBRARY_ROOT = BASE_DIR / "Tender_Aligned_FinalLibrary"

# Map of categories to their folder names in the repo
CATEGORY_FOLDERS = {
    "Corporate Profile": LIBRARY_ROOT / "Corporate Profile",
    "Technical Profile": LIBRARY_ROOT / "Technical Profile",
    "Security Profile": LIBRARY_ROOT / "Security Profile",
    "Services And Delivery": LIBRARY_ROOT / "Services And Delivery",
}

# ---------- HELPER: READ WORD DOCUMENT ----------
def load_docx_text(file_path: Path) -> str:
    if not file_path.exists():
        return "❗ This document does not exist: " + str(file_path)
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n\n".join(p for p in paragraphs if p.strip())

def list_docx_files(folder: Path):
    if not folder.exists():
        return []
    return sorted([p for p in folder.glob("*.docx")])

# ---------- SIMPLE INDEXING STUB ----------
def index_library():
    """
    Very simple in-memory indexing stub.
    - Walks all category folders
    - Reads each .docx
    - Stores text chunks in st.session_state['library_index']
    This is where you can later plug in embeddings / vector DB, etc.
    """
    index = []
    for category, folder in CATEGORY_FOLDERS.items():
        doc_files = list_docx_files(folder)
        for doc_path in doc_files:
            text = load_docx_text(doc_path)
            if not text.strip():
                continue

            # For now, treat whole document as a single chunk.
            # Later you can split into smaller chunks if needed.
            index.append({
                "category": category,
                "file_name": doc_path.name,
                "file_path": str(doc_path),
                "text": text,
            })

    st.session_state["library_index"] = index
    return index

# ---------- SIMPLE RETRIEVAL: LIBRARY ENTRIES RELEVANT TO A CLAUSE ----------
def get_relevant_library_entries(clause_text: str, top_k: int = 5):
    """
    Naive relevance scoring:
    - Lowercase clause text
    - Remove very short words
    - Score each library entry by how many keywords it contains
    Returns top_k entries sorted by score.
    """
    library_index = st.session_state.get("library_index", [])
    if not library_index:
        return []

    # Basic keyword set from clause
    words = re.findall(r"\w+", clause_text.lower())
    keywords = {w for w in words if len(w) > 3}  # remove tiny words: the, and, etc.

    scored = []
    for entry in library_index:
        text_lower = entry["text"].lower()
        score = sum(1 for w in keywords if w in text_lower)
        if score > 0:
            scored.append((score, entry))

    scored.sort(key=lambda x: x[0], reverse=True)
    top_entries = [e for score, e in scored[:top_k]]
    return top_entries

# ---------- OPENAI HELPERS ----------
def get_openai_api_key():
    """
    Read API key from Streamlit secrets.
    Supports:
      OPENAI_API_KEY = "sk-..."
    or:
      [openai]
      api_key = "sk-..."
    """
    api_key = st.secrets.get("OPENAI_API_KEY")
    if not api_key and "openai" in st.secrets:
        api_key = st.secrets["openai"].get("api_key")

    if not api_key:
        st.error("❌ OPENAI_API_KEY not found in secrets. Please add it in Streamlit Cloud settings.")
        return None
    return api_key


def generate_openai_response(prompt_text: str):
    """
    Call OpenAI Chat Completions API (new SDK) and return the response text.
    Uses `OpenAI` client and `client.chat.completions.create`.
    """
    if not OPENAI_AVAILABLE or OpenAI is None:
        st.error("❌ OpenAI Python client is not installed. Check requirements.txt for 'openai'.")
        return None

    api_key = get_openai_api_key()
    if api_key is None:
        return None

    try:
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4.1-mini",   # adjust if you want another model
            messages=[{"role": "user", "content": prompt_text}],
            temperature=0.2,
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"❌ OpenAI API call failed: {e}")
        return None

# ---------- HELPERS FOR TENDER UPLOAD / CLAUSE EXTRACTION ----------
def extract_text_from_pdf(uploaded_file) -> str:
    """Extract plain text from a PDF file using pdfplumber."""
    if not PDF_SUPPORT or pdfplumber is None:
        st.error("PDF support is not available on this deployment. Please upload DOCX or Excel instead.")
        return ""
    all_text = []
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                all_text.append(page_text)
    return "\n".join(all_text)

def extract_text_from_docx_file(uploaded_file) -> str:
    """Extract plain text from an uploaded DOCX file."""
    doc = Document(uploaded_file)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(p for p in paragraphs if p.strip())

def extract_text_from_excel_file(uploaded_file) -> str:
    """Extract text from Excel (.xlsx or .xls) by concatenating non-empty cell values."""
    try:
        # Read all sheets, no headers (tender text might be anywhere)
        sheets = pd.read_excel(uploaded_file, sheet_name=None, header=None)
    except Exception as e:
        st.error(f"❌ Could not read Excel file: {e}")
        return ""

    parts = []
    for sheet_name, df in sheets.items():
        parts.append(f"Sheet: {sheet_name}")
        for _, row in df.iterrows():
            for cell in row:
                if pd.notna(cell):
                    parts.append(str(cell))
    return "\n".join(parts)

def extract_clauses_from_text(raw_text: str):
    """
    Extract clauses based on patterns like:
    3.14, 3.15, 4.2.1 at the start of lines.

    Returns a list of dicts:
    [{"clause_no": "3.14", "clause_text": "..."}, ...]
    """
    lines = [line.strip() for line in raw_text.splitlines()]
    clause_pattern = re.compile(r'^(\d+(?:\.\d+)*)\s+(.*)$')

    clauses = []
    current_clause = None

    for line in lines:
        if not line:
            continue

        match = clause_pattern.match(line)
        if match:
            # Start of a new clause
            if current_clause:
                clauses.append(current_clause)

            clause_no = match.group(1)
            clause_text = match.group(2).strip()
            current_clause = {
                "clause_no": clause_no,
                "clause_text": clause_text
            }
        else:
            # Continuation of the previous clause
            if current_clause:
                current_clause["clause_text"] += " " + line

    # Append the last clause if present
    if current_clause:
        clauses.append(current_clause)

    return clauses

# ---------- SIDEBAR NAVIGATION ----------
st.sidebar.header("🧭 Navigation")
page = st.sidebar.radio(
    "Go to:",
    [
        "📖 View Documents",
        "🧠 Prepare / Index Library",
        "📄 Upload Tender & Extract Clauses",
        "📝 Generate Responses",
    ],
    index=0
)

# ---------- PAGE: PREPARE / INDEX LIBRARY ----------
if page == "🧠 Prepare / Index Library":
    st.subheader("🧠 Prepare / Index Wavetec Library")

    st.markdown(
        """
        This page prepares your Wavetec Tender Library for automated RFP responses.

        **What this button will (currently) do:**
        - Scan all categories and `.docx` files
        - Read their full text
        - Build an in-memory index in `st.session_state["library_index"]`
        - Show you how many documents were indexed

        Later, this function can be extended to:
        - Split documents into smaller chunks
        - Generate embeddings
        - Store them in a vector database
        """
    )

    if st.button("🚀 Index / Refresh Library"):
        with st.spinner("Indexing library..."):
            index = index_library()
        st.success(f"✅ Indexed {len(index)} document entries into memory.")

        # Optional: small preview table (file + category)
        if index:
            st.markdown("### 📄 Indexed Documents (Preview)")
            preview = [
                {"Category": item["category"], "File": item["file_name"]}
                for item in index
            ]
            st.dataframe(preview)

# ---------- PAGE: UPLOAD TENDER & EXTRACT CLAUSES ----------
elif page == "📄 Upload Tender & Extract Clauses":
    st.subheader("📄 Upload Tender & Extract Clauses")

    st.markdown(
        """
        Upload a tender file.

        The app will:
        - Extract the full text
        - Detect clause numbers like `3.14`, `3.15`, `4.2.1` at the start of lines
        - Build a **Tender Response Map** (clause number + clause text)
        - Store it in `st.session_state["tender_clauses"]` for later use.
        """
    )

    # Allowed file types depend on whether PDF support is available
    base_types = ["docx", "xlsx", "xls"]
    if PDF_SUPPORT:
        allowed_types = base_types + ["pdf"]
        st.info("📎 You can upload PDF, Word (.docx), or Excel (.xlsx/.xls) tenders.")
    else:
        allowed_types = base_types
        st.warning("⚠️ PDF support is not available on this deployment. Please upload Word (.docx) or Excel (.xlsx/.xls).")

    uploaded_file = st.file_uploader(
        "Upload Tender Document",
        type=allowed_types,
        help="Accepted formats: " + ", ".join(f".{ext}" for ext in allowed_types)
    )

    if uploaded_file is not None:
        st.info(f"📁 File uploaded: **{uploaded_file.name}**")

        if st.button("🔍 Extract Clauses"):
            with st.spinner("Extracting text and detecting clauses..."):
                filename = uploaded_file.name.lower()

                # Step 1: Extract raw text depending on extension
                if filename.endswith(".pdf"):
                    raw_text = extract_text_from_pdf(uploaded_file)
                elif filename.endswith(".docx"):
                    raw_text = extract_text_from_docx_file(uploaded_file)
                elif filename.endswith(".xlsx") or filename.endswith(".xls"):
                    raw_text = extract_text_from_excel_file(uploaded_file)
                else:
                    st.error("❌ Unsupported file type.")
                    raw_text = ""

                if not raw_text or not raw_text.strip():
                    st.error("❌ No text could be extracted from this file.")
                else:
                    # Step 2: Extract clauses
                    clauses = extract_clauses_from_text(raw_text)

                    if not clauses:
                        st.warning(
                            "⚠️ No clauses were detected. "
                            "Check if the document uses standard numbering like '3.14', '4.2.1' at the start of lines."
                        )
                    else:
                        # Save in session_state for later steps (response generation)
                        st.session_state["tender_clauses"] = clauses

                        st.success(f"✅ Extracted {len(clauses)} clauses from the tender.")

                        # Show a preview table
                        df_clauses = pd.DataFrame(clauses)
                        st.markdown("### 📋 Tender Response Map (Preview)")
                        st.dataframe(df_clauses, use_container_width=True)

                        with st.expander("🔎 View first 5 clauses (full text)"):
                            for row in clauses[:5]:
                                st.markdown(f"**Clause {row['clause_no']}**")
                                st.write(row["clause_text"])
                                st.markdown("---")
    else:
        st.info("📥 Please upload a tender file to begin.")

# ---------- PAGE: 📝 GENERATE RESPONSES ----------
elif page == "📝 Generate Responses":
    st.subheader("📝 Generate Clause-by-Clause Responses")

    st.markdown(
        """
        This page helps you generate Wavetec responses for each tender clause.

        **Workflow:**
        1. Make sure you have:
           - Indexed the Wavetec library on the **"🧠 Prepare / Index Library"** page.
           - Extracted clauses on the **"📄 Upload Tender & Extract Clauses"** page.
        2. Select a clause from the dropdown.
        3. The app will find the most relevant library entries (simple keyword match for now).
        4. You can:
           - Copy the prepared prompt into ChatGPT **or**
           - Click **"🤖 Generate with OpenAI"** to call the API directly.
        """
    )

    tender_clauses = st.session_state.get("tender_clauses")
    library_index = st.session_state.get("library_index")

    if not tender_clauses:
        st.warning("⚠️ No tender clauses found. Please first upload a tender and run **Extract Clauses**.")
    elif not library_index:
        st.warning("⚠️ Library index is empty. Please go to **🧠 Prepare / Index Library** and index your documents first.")
    else:
        # Build a small display list for the selectbox
        options = []
        for c in tender_clauses:
            snippet = c["clause_text"][:80].replace("\n", " ")
            options.append(f"{c['clause_no']} – {snippet}...")

        selected_label = st.selectbox("Select a clause to respond to:", options)

        # Find the underlying clause dict
        selected_index = options.index(selected_label)
        selected_clause = tender_clauses[selected_index]

        st.markdown("### 📌 Selected Clause")
        
        st.markdown(f"**Clause {selected_clause['clause_no']}**")
        st.write(selected_clause["clause_text"])

        # Get relevant library entries
        with st.spinner("Finding relevant Wavetec library content..."):
            relevant_entries = get_relevant_library_entries(selected_clause["clause_text"], top_k=5)

        if not relevant_entries:
            st.warning("No relevant library entries were found with the simple keyword search. You can still draft a response manually.")
        else:
            st.markdown("### 📚 Top Matching Library Entries (Preview)")
            for i, entry in enumerate(relevant_entries, start=1):
                st.markdown(f"**{i}. {entry['category']} – {entry['file_name']}**")
                # Show only a small snippet
                snippet = entry["text"][:400].replace("\n", " ")
                st.write(snippet + "...")
                st.markdown("---")

            # Build a compiled context string
            context_blocks = []
            for entry in relevant_entries:
                block = (
                    f"Category: {entry['category']}\n"
                    f"File: {entry['file_name']}\n"
                    f"Content:\n{entry['text']}\n"
                    "-------------------------\n"
                )
                context_blocks.append(block)

            compiled_context = "\n".join(context_blocks)

            # Prepare a ready-made prompt for ChatGPT / API
            prompt_text = f"""You are the Wavetec RFP Response Engine.

Use ONLY the library context below to answer the clause. Do NOT invent facts that are not supported by the context.

Clause number: {selected_clause['clause_no']}
Clause text:
\"\"\"{selected_clause['clause_text']}\"\"\"


Wavetec Library Context:
\"\"\"{compiled_context}\"\"\"


TASK:
1. Write a detailed, structured, bid-winning response to this clause from Wavetec's perspective.
2. Use a formal, government/RFP-appropriate tone.
3. Include all relevant technical, architectural, security, delivery, and operational details you can find in the context.
4. After the main response, add:

**Gaps / Missing Information:**
- Bullet list of items where the library does not contain enough detail.

**Assumptions:**
- Bullet list of assumptions you are making to answer this clause.

Return the answer in markdown format.
"""

            st.markdown("### ✍️ Prepared Prompt (you can still copy & paste if needed)")
            st.text_area(
                "Prompt for ChatGPT / LLM",
                value=prompt_text,
                height=300
            )

            # --- Direct OpenAI API call button ---
            if st.button("🤖 Generate with OpenAI"):
                with st.spinner("Calling OpenAI to generate the response..."):
                    answer = generate_openai_response(prompt_text)

                if answer:
                    st.markdown("### ✅ Generated Response")
                    st.markdown(answer)

            st.info("You can either copy the prompt above or use the '🤖 Generate with OpenAI' button for one-click generation.")

# ---------- PAGE: VIEW DOCUMENTS (EXISTING BEHAVIOUR) ----------
elif page == "📖 View Documents":
    st.sidebar.header("📂 Document Library")

    category = st.sidebar.selectbox(
        "Select a category:",
        list(CATEGORY_FOLDERS.keys())
    )

    folder_path = CATEGORY_FOLDERS[category]
    doc_files = list_docx_files(folder_path)

    if not doc_files:
        st.warning(f"No .docx files found in folder: {folder_path.name}")
    else:
        doc_display_names = [f.name for f in doc_files]
        selected_doc_name = st.sidebar.selectbox(
            "Select a document:",
            doc_display_names
        )

        selected_doc_path = folder_path / selected_doc_name

        st.subheader(f"{category} → {selected_doc_name}")
        st.caption(f"Source file: {selected_doc_path.relative_to(BASE_DIR)}")

        content = load_docx_text(selected_doc_path)
        st.markdown(content)

::contentReference[oaicite:0]{index=0}
